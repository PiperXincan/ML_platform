from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

OUT = '王心粲_策略算法与数据科学家简历.docx'

NAVY = '17365D'
BLUE = '2E74B5'
TEXT = '1F2937'
MUTED = '5B6472'
LIGHT = 'EAF1F8'

def set_east_asia(run, font='Microsoft YaHei'):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement('w:rFonts')
        rpr.append(fonts)
    fonts.set(qn('w:eastAsia'), font)
    fonts.set(qn('w:ascii'), 'Aptos')
    fonts.set(qn('w:hAnsi'), 'Aptos')

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(width / 2.54 * 1440)))
            tcW.set(qn('w:type'), 'dxa')

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ['val', 'sz', 'color', 'space']:
                if key in kwargs[edge]:
                    element.set(qn('w:{}'.format(key)), str(kwargs[edge][key]))

def add_run(p, text, size=9.2, bold=False, color=TEXT, font='Microsoft YaHei'):
    r = p.add_run(text)
    set_east_asia(r, font)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return r

def add_para(doc, text='', size=9.2, after=1, before=0, line=1.08, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.left_indent = Cm(indent)
    if text:
        add_run(p, text, size=size)
    return p

def section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = add_run(p, title, size=10.8, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '9')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), BLUE)
    borders.append(bottom)
    pPr.append(borders)
    return p

def role_header(doc, company, role, dates):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(t, [10.7, 6.0])
    for cell in t.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        set_cell_border(cell, top={'val':'nil'}, bottom={'val':'nil'}, left={'val':'nil'}, right={'val':'nil'})
    p = t.cell(0,0).paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_run(p, company, size=9.7, bold=True, color=TEXT)
    add_run(p, '  |  ' + role, size=9.2, color=MUTED)
    p2 = t.cell(0,1).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(1)
    add_run(p2, dates, size=8.8, color=MUTED)
    return t

def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.36)
    p.paragraph_format.first_line_indent = Cm(-0.30)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, '• ', size=9.0, color=BLUE)
    add_run(p, text, size=9.0)
    return p

def education(doc, school, degree, dates, detail):
    role_header(doc, school, degree, dates)
    p = add_para(doc, after=1)
    add_run(p, detail, size=8.8, color=MUTED)

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.05)
sec.bottom_margin = Cm(1.05)
sec.left_margin = Cm(1.15)
sec.right_margin = Cm(1.15)
sec.header_distance = Cm(0.5)
sec.footer_distance = Cm(0.5)

styles = doc.styles
styles['Normal'].font.name = 'Microsoft YaHei'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
styles['Normal'].font.size = Pt(9.2)

# Header
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(1)
add_run(p, '王心粲', size=20, bold=True, color=NAVY)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(5)
add_run(p2, '策略算法工程师 / 数据科学家', size=10.5, bold=True, color=BLUE)
add_run(p2, '    25岁  |  中共党员  |  18981934809  |  wangxc_piper@163.com', size=8.8, color=MUTED)

# Summary callout
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_widths(t, [16.7])
c = t.cell(0,0)
set_cell_shading(c, LIGHT)
set_cell_margins(c, top=95, bottom=95, start=140, end=140)
set_cell_border(c, top={'val':'single','sz':'6','color':BLUE}, bottom={'val':'single','sz':'6','color':BLUE}, left={'val':'single','sz':'6','color':BLUE}, right={'val':'single','sz':'6','color':BLUE})
p = c.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
add_run(p, '个人概述  ', size=9.2, bold=True, color=NAVY)
add_run(p, '金融科技硕士，具备端到端预测建模、用户分层、文本挖掘、图特征与因果推断实践。已完成10+模型开发及迭代，擅长将业务问题抽象为可验证的数据策略，目标岗位为策略算法、数据科学、增长/营销智能化。', size=8.9)

section_title(doc, '核心能力')
p = add_para(doc, after=1, line=1.08)
add_run(p, '建模与策略：', size=8.9, bold=True, color=NAVY)
add_run(p, '特征工程、LGBM/XGBoost、模型融合、Uplift Model、图谱建模、异常检测、模型监控与归因分析', size=8.9)
p = add_para(doc, after=1, line=1.08)
add_run(p, '数据与工程：', size=8.9, bold=True, color=NAVY)
add_run(p, 'Python（NumPy、Pandas、Matplotlib、Sklearn）、SQL、DataWorks、天玑、DolphinScheduler、BI看板', size=8.9)
p = add_para(doc, after=1, line=1.08)
add_run(p, 'AI与业务：', size=8.9, bold=True, color=NAVY)
add_run(p, 'Transformer/CNN文本建模、AI Agent使用与竞品分析、业务需求分析、跨团队协作与项目测试', size=8.9)

section_title(doc, '工作经历')
role_header(doc, '奇富科技', '算法工程师', '2026.01 - 2026.03')
bullet(doc, '主导贷前、贷中、贷后等多场景预测模型的全流程开发与迭代，累计完成10+模型：从业务分析、特征工程、LGBM/XGBoost训练，到上线监控与策略落地。')
bullet(doc, '构建多数据源融合模型，量化各数据源的增量价值并完成消融实验，使核心模型KS提升3%。')
bullet(doc, '从0到1搭建手机号-地址关联图谱并提取邻居聚合特征，提升存量用户风险识别精度，并缓解新用户冷启动问题。')
bullet(doc, '构建差分响应（Uplift）模型评估策略净效应，识别高敏感人群，为差异化资源分配提供量化依据。')

role_header(doc, '桔子数科', '风控数据分析师', '2025.01 - 2025.12')
bullet(doc, '使用CNN与Transformer对历史外呼文本及客户行为建模，预测复播后的动支概率，支持营销名单筛选与转化效率优化。')
bullet(doc, '独立完成新增外部数据源的全流程评估与定制化联合建模：对标基准评分、输出评估报告和模型方案，辅助数据采购与策略决策。')
bullet(doc, '参与多数据源模型融合与消融实验，拆解不同数据源的边际贡献，沉淀可复用的评估方法。')

role_header(doc, '上海恒生聚源数据服务有限公司', '金融市场产品部实习生', '2024.06 - 2024.09')
bullet(doc, '负责大模型理财助手的竞品研究、交互需求分析和产品测试；参考主流智能投顾与问答产品，输出需求分析文档并跟进项目推进。')
bullet(doc, '完成Agent智能体平台竞品分析，覆盖Coze、文心智能体平台、agentUniverse等产品，为AI应用方案设计提供输入。')

section_title(doc, '代表项目')
role_header(doc, '模型监控与业务归因分析', '数据科学 / 策略分析', '工作项目')
bullet(doc, '参与搭建模型与业务监控看板；针对入催率异常上涨，完成客群结构变化与单客群风险上升的归因拆解，支持策略及时调整。')
role_header(doc, '差异化额度策略模型', '预测建模 / 精细化运营', '工作项目')
bullet(doc, '面向高、中、低额度客群建立差异化模型，并结合联合分布与逾期率交叉分析，为精细化提额策略提供数据依据。')

section_title(doc, '教育背景与荣誉')
education(doc, '华东理工大学', '金融（金融科技方向）硕士', '2023.09 - 2025.06', '专业排名第1；课程：机器学习与金融应用、NLP与知识图谱、金融科技与风险管理。荣誉：研究生金融科技创新大赛校内一等奖、优秀毕业生、一等学业奖学金。')
education(doc, '南京林业大学', '金融工程学士', '2019.09 - 2023.06', '专业排名前3%；荣誉：MathorCup数学建模国赛二等奖、美国大学生数学建模比赛S奖。')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0
add_run(p, '其他：', size=8.4, bold=True, color=NAVY)
add_run(p, '英语CET-4/6，具备中英文读写能力；熟练使用Excel、Word、PPT，能够完成数据分析报告与业务展示。', size=8.4, color=MUTED)

# Footer
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_after = Pt(0)
add_run(footer, '王心粲 | 策略算法工程师 / 数据科学家', size=7.5, color=MUTED)

doc.core_properties.author = '王心粲'
doc.core_properties.title = '王心粲 - 策略算法与数据科学家简历'
doc.save(OUT)
print(OUT)
